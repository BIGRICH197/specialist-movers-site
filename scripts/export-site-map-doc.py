"""Export site map to Downloads as .docx and .html."""
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOWNLOADS = Path.home() / "Downloads"
DOMAIN = "https://specialistmovers.co.nz"

CITY_URLS = [
    ("/piano-movers/auckland", "/piano-movers/hamilton"),
    ("/services/house-moving/auckland", "/services/house-moving/hamilton"),
    ("/services/office-moving/auckland", "/services/office-moving/hamilton"),
    ("/services/commercial-moving/auckland", "/services/commercial-moving/hamilton"),
    ("/services/packing-services/auckland", "/services/packing-services/hamilton"),
    ("/services/hard-to-shift/auckland", "/services/hard-to-shift/hamilton"),
    ("/services/cleaning-services/auckland", "/services/cleaning-services/hamilton"),
    ("/services/international-moving/auckland", "/services/international-moving/hamilton"),
    ("/services/loading-unloading/auckland", "/services/loading-unloading/hamilton"),
    ("/services/winz-quotes/auckland", "/services/winz-quotes/hamilton"),
    ("/services/storage/auckland", "/services/storage/hamilton"),
    ("/services/moving/auckland", "/services/moving/hamilton"),
]

CITY_SERVICES = [
    ("Piano", *CITY_URLS[0]),
    ("House moving", *CITY_URLS[1]),
    ("Office moving", *CITY_URLS[2]),
    ("Commercial moving", *CITY_URLS[3]),
    ("Packing", *CITY_URLS[4]),
    ("Hard to shift", *CITY_URLS[5]),
    ("Exit cleaning", *CITY_URLS[6]),
    ("International", *CITY_URLS[7]),
    ("Loading / unloading", *CITY_URLS[8]),
    ("WINZ quote", *CITY_URLS[9]),
    ("Moving storage", *CITY_URLS[10]),
    ("Moving hub", *CITY_URLS[11]),
]


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def build_docx(path: Path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Specialist Movers — Full site map", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph("Site: specialistmovers.co.nz (Next.js on Vercel)")
    doc.add_paragraph("Generated: 3 June 2026")
    doc.add_paragraph("Total public pages: ~159")
    doc.add_paragraph(
        "Use this document for SEO backlinks, internal linking, and redirects from the old WordPress sites."
    )

    add_heading(doc, "Quick counts", 2)
    table = doc.add_table(rows=8, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Core marketing", "11"),
        ("Piano hub + types + cities", "7"),
        ("Services (hubs + city landings)", "27"),
        ("Moving / storage clusters", "8"),
        ("Locations (areas)", "95"),
        ("Blog", "4"),
        ("Internal / dev only", "3"),
        ("City SEO landings (backlinks)", "24"),
    ]
    for i, (a, b) in enumerate(rows):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b

    add_heading(doc, "Site structure (overview)", 2)
    tree = """/  Homepage
├── /about, /why-us, /faq, /reviews, /contact, /policies
├── /blog (+ 3 posts)
├── /services (hub)
│   ├── house-moving, office-moving, commercial-moving (+ Auckland & Hamilton each)
│   ├── packing-services, hard-to-shift, cleaning-services (+ cities)
│   ├── international-moving, loading-unloading, winz-quotes (+ cities)
│   ├── /moving (hub + cities) → local-moving, regional-moving
│   └── /storage (hub + cities) → short/long-term, in-transit, overnight
├── /piano-movers (hub + Auckland & Hamilton + 4 piano types)
└── /locations (hub + 94 area pages)"""
    doc.add_paragraph(tree)

    add_heading(doc, "City landing pages — backlink targets (24)", 2)
    doc.add_paragraph(
        "Use these full URLs for targeted backlinks. Prepend: " + DOMAIN
    )
    ct = doc.add_table(rows=1 + len(CITY_SERVICES), cols=3)
    ct.style = "Table Grid"
    hdr = ct.rows[0].cells
    hdr[0].text = "Service"
    hdr[1].text = "Auckland"
    hdr[2].text = "Hamilton"
    for i, (name, akl, ham) in enumerate(CITY_SERVICES, start=1):
        ct.rows[i].cells[0].text = name
        ct.rows[i].cells[1].text = DOMAIN + akl
        ct.rows[i].cells[2].text = DOMAIN + ham

    add_heading(doc, "Core pages", 2)
    core = [
        ("/", "Homepage"),
        ("/about", "About"),
        ("/why-us", "Why choose us"),
        ("/faq", "FAQs"),
        ("/reviews", "Google reviews"),
        ("/contact", "Contact + quote"),
        ("/policies", "Policies"),
        ("/services", "Services hub"),
        ("/locations", "Locations hub"),
        ("/blog", "Blog hub"),
        ("/piano-movers", "Piano hub"),
    ]
    t = doc.add_table(rows=len(core), cols=2)
    t.style = "Table Grid"
    for i, (url, purpose) in enumerate(core):
        t.rows[i].cells[0].text = url
        t.rows[i].cells[1].text = purpose

    add_heading(doc, "Redirects (301)", 2)
    redirects = [
        ("/services/piano-movers", "/piano-movers"),
        ("/services/piano-movers/auckland", "/piano-movers/auckland"),
        ("/services/piano-movers/hamilton", "/piano-movers/hamilton"),
        ("/services/moving/international-moving", "/services/international-moving"),
        ("/services/storage/piano-storage", "/piano-movers/piano-storage"),
    ]
    rt = doc.add_table(rows=len(redirects), cols=2)
    rt.style = "Table Grid"
    for i, (f, t_) in enumerate(redirects):
        rt.rows[i].cells[0].text = f
        rt.rows[i].cells[1].text = t_

    add_heading(doc, "Locations", 2)
    doc.add_paragraph(
        "/locations — hub plus 94 pages (/locations/{suburb-or-town}). "
        "Use city service pages for backlinks; location pages for suburb long-tail."
    )

    add_heading(doc, "Blog", 2)
    for slug in [
        "the-ultimate-guide-to-house-moving-in-auckland",
        "diy-packing-vs-professional-packing-services",
        "stress-free-moving-in-auckland-expert-tips-from-specialist-movers",
    ]:
        doc.add_paragraph(f"/blog/{slug}", style="List Bullet")

    doc.add_paragraph("")
    doc.add_paragraph("Specialist Movers · Auckland & Hamilton bases · KB Logistics Limited")

    doc.save(path)


def build_html(path: Path):
    city_rows = "\n".join(
        f"<tr><td>{name}</td><td><a href=\"{DOMAIN}{a}\">{DOMAIN}{a}</a></td>"
        f"<td><a href=\"{DOMAIN}{h}\">{DOMAIN}{h}</a></td></tr>"
        for name, a, h in CITY_SERVICES
    )
    path_block = "\n".join(f"{a}\n{h}" for a, h in CITY_URLS)
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <title>Specialist Movers Site Map</title>
  <style>
    body {{ font-family: Calibri, Segoe UI, Arial, sans-serif; max-width: 900px; margin: 2rem auto; padding: 0 1rem; color: #1a0a2e; }}
    h1 {{ color: #5b2c6f; }}
    h2 {{ color: #5b2c6f; margin-top: 2rem; border-bottom: 2px solid #f4c430; padding-bottom: 0.25rem; }}
    table {{ border-collapse: collapse; width: 100%; margin: 1rem 0; }}
    th, td {{ border: 1px solid #ccc; padding: 8px 10px; text-align: left; }}
    th {{ background: #5b2c6f; color: white; }}
    tr:nth-child(even) {{ background: #f8f4fa; }}
    pre {{ background: #f4f4f4; padding: 1rem; overflow-x: auto; font-size: 13px; }}
    a {{ color: #5b2c6f; }}
    .meta {{ color: #555; }}
  </style>
</head>
<body>
  <h1>Specialist Movers — Full site map</h1>
  <p class="meta"><strong>Site:</strong> specialistmovers.co.nz · <strong>Generated:</strong> 3 June 2026 · <strong>~159 pages</strong></p>
  <p>Use for SEO backlinks, internal linking, and WordPress redirects.</p>

  <h2>Quick counts</h2>
  <table>
    <tr><th>Section</th><th>Pages</th></tr>
    <tr><td>Core marketing</td><td>11</td></tr>
    <tr><td>Piano + cities</td><td>7</td></tr>
    <tr><td>Services + cities</td><td>27</td></tr>
    <tr><td>Moving / storage clusters</td><td>8</td></tr>
    <tr><td>Locations</td><td>95</td></tr>
    <tr><td>Blog</td><td>4</td></tr>
    <tr><td><strong>City SEO landings</strong></td><td><strong>24</strong></td></tr>
  </table>

  <h2>City pages — backlink targets (24)</h2>
  <table>
    <tr><th>Service</th><th>Auckland</th><th>Hamilton</th></tr>
    {city_rows}
  </table>

  <h2>Site tree</h2>
  <pre>/  Homepage
├── /about, /why-us, /faq, /reviews, /contact, /policies
├── /blog (+ 3 posts)
├── /services (+ each service has /auckland and /hamilton)
├── /piano-movers (+ /auckland, /hamilton, 4 piano types)
└── /locations (+ 94 suburbs/towns)</pre>

  <h2>All 24 city paths (copy/paste)</h2>
  <pre>{path_block}</pre>
  <p>Prepend: <strong>{DOMAIN}</strong></p>

  <h2>Redirects</h2>
  <table>
    <tr><th>From</th><th>To</th></tr>
    <tr><td>/services/piano-movers</td><td>/piano-movers</td></tr>
    <tr><td>/services/piano-movers/auckland</td><td>/piano-movers/auckland</td></tr>
    <tr><td>/services/piano-movers/hamilton</td><td>/piano-movers/hamilton</td></tr>
  </table>

  <p><em>Specialist Movers · Auckland &amp; Hamilton bases</em></p>
</body>
</html>"""
    path.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    DOWNLOADS.mkdir(parents=True, exist_ok=True)
    docx_path = DOWNLOADS / "00-Specialist-Movers-Site-Map.docx"
    html_path = DOWNLOADS / "00-Specialist-Movers-Site-Map.html"
    build_docx(docx_path)
    build_html(html_path)
    print(f"Wrote {docx_path}")
    print(f"Wrote {html_path}")
